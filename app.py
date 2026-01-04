import streamlit as st
import gspread
from oauth2client.service_account import ServiceAccountCredentials
from datetime import datetime, timedelta, date
import io
import smtplib
from email.mime.text import MIMEText
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import time
import hashlib

# =========================================================
# 0) 基礎設定
# =========================================================
SHEET_URL = "https://docs.google.com/spreadsheets/d/1zXHavJqhOBq1-m_VR7sxMkeOHdXoD9EmQCEM1Nl816I/edit?usp=sharing"

PROVIDER_NAME = "高如慧"
BANK_NAME = "中國信託商業銀行"
BANK_CODE = "822"
ACCOUNT_NUMBER = "783540208870"
REMOTE_SUPPORT_URL = "https://remotedesktop.google.com/support"
CREATIVES_UPLOAD_URL = "https://metaads-dtwbm3ntmprhjvpv6ptmec.streamlit.app/"
BM_TUTORIAL_URL = "https://www.youtube.com/watch?v=caoZAO8tyNs"

st.set_page_config(
    page_title="廣告投放服務｜合約＋啟動資料收集",
    page_icon="📝",
    layout="centered"
)

# =========================================================
# 1) 工具函式
# =========================================================
def make_hash(password):
    return hashlib.sha256(password.encode()).hexdigest()

def check_password(input_pw, db_pw):
    if len(db_pw) == 64:
        return make_hash(input_pw) == db_pw
    return input_pw == db_pw

@st.cache_resource
def get_gsheet_client():
    scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
    creds_dict = dict(st.secrets["gcp_service_account"])
    creds = ServiceAccountCredentials.from_json_keyfile_dict(creds_dict, scope)
    client = gspread.authorize(creds)
    return client

def get_worksheet():
    client = get_gsheet_client()
    sheet = client.open_by_url(SHEET_URL)
    return sheet.get_worksheet(0)

def send_email(subject, body):
    try:
        sender = st.secrets["email"]["sender_email"]
        password = st.secrets["email"]["sender_password"]
        receiver = st.secrets["email"]["receiver_email"]
        msg = MIMEText(body, 'plain', 'utf-8')
        msg['Subject'] = subject
        msg['From'] = sender
        msg['To'] = receiver
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
            server.login(sender, password)
            server.send_message(msg)
        return True
    except Exception as e:
        st.error(f"Email 發送失敗: {e}")
        return False

# =========================================================
# 2) 資料處理邏輯 (嚴格對應 CSV 順序)
# =========================================================
def find_user_row(email):
    ws = get_worksheet()
    records = ws.get_all_records()
    for i, record in enumerate(records):
        if record.get("Email") == email:
            return i + 2, record
    return None, None

def save_phase1_new(data_dict):
    ws = get_worksheet()
    def s(key): return data_dict.get(key, "")
    hashed_default = make_hash("dennis")
    # Email(1), case_id(2), party_a(3), provider(4), plan(5), start_date(6), pay_day(7), pay_date(8)
    row = [
        s("Email"), s("case_id"), s("party_a"), PROVIDER_NAME, s("plan"), 
        str(s("start_date")), s("pay_day"), str(s("pay_date")) if s("pay_date") else "",
        "FALSE", "FALSE", "FALSE", "FALSE", # chk boxes (9-12)
        "", "", # URLs (13-14)
        "", "", "", # Comps (15-17)
        "", "", "", "", # Problems & Budget (18-21)
        datetime.now().strftime("%Y-%m-%d %H:%M:%S"), # last_update (22)
        "contract", s("plan"), f"{s('case_id')} ({s('party_a')})", # 23-25
        "FALSE", "FALSE", hashed_default # 26-28
    ]
    ws.append_row(row)

# --- [新增] 用於登入後更新 Phase 1 方案資料的函式 ---
def update_phase1(row_num, plan, start_date, pay_day, pay_date):
    ws = get_worksheet()
    # Sheet 欄位順序: plan(5), start_date(6), pay_day(7), pay_date(8)
    # 注意: update_cell 是 (row, col, value)
    ws.update_cell(row_num, 5, plan)
    ws.update_cell(row_num, 6, str(start_date))
    ws.update_cell(row_num, 7, pay_day)
    ws.update_cell(row_num, 8, str(pay_date) if pay_date else "")
    # 同步更新後面用來做合約紀錄的欄位 (第 24 欄也是 plan)
    ws.update_cell(row_num, 24, plan)

def update_phase2(row_num, p2_data):
    ws = get_worksheet()
    cells = []
    def Cell(col, val): return gspread.Cell(row_num, col, str(val))
    # 勾選框
    cells.append(Cell(9, p2_data["chk_ad_account"]))
    cells.append(Cell(10, p2_data["chk_pixel"]))
    cells.append(Cell(11, p2_data["chk_fanpage"]))
    cells.append(Cell(12, p2_data["chk_bm"]))
    # 網址
    cells.append(Cell(13, p2_data["fanpage_url"]))
    cells.append(Cell(14, p2_data["landing_url"]))
    # 競品
    cells.append(Cell(15, p2_data["comp1"]))
    cells.append(Cell(16, p2_data["comp2"]))
    cells.append(Cell(17, p2_data["comp3"]))
    # 策略問題
    cells.append(Cell(18, p2_data["who_problem"]))
    cells.append(Cell(19, p2_data["what_problem"]))
    cells.append(Cell(20, p2_data["how_solve"]))
    cells.append(Cell(21, p2_data["budget"]))
    # 系統資訊
    cells.append(Cell(22, datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
    cells.append(Cell(26, p2_data["chk_remote"]))
    cells.append(Cell(27, p2_data["chk_creatives"]))
    ws.update_cells(cells)

def update_password(row_num, new_pw):
    ws = get_worksheet()
    ws.update_cell(row_num, 28, make_hash(new_pw))

# =========================================================
# 3) Word 合約生成 (完整 14 條款版)
# =========================================================
def set_run_font(run, size=10.5, bold=False):
    run.font.name = "Microsoft JhengHei"
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

def generate_docx_bytes(party_a, email, payment_opt, start_dt, pay_day, pay_dt, case_num):
    doc = Document()
    
    # 設定頁邊距
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # 設定預設字體與行距
    style = doc.styles['Normal']
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)

    # 1. 標題
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("廣告投放服務合約書")
    set_run_font(run, size=18, bold=True)
    
    if case_num:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_sub = sub.add_run(f"案件編號：{case_num}")
        set_run_font(run_sub, size=10)

    # 2. 前言
    doc.add_paragraph("")
    p_intro = doc.add_paragraph()
    run_intro = p_intro.add_run(f"立合約書人：委託人 {party_a}（以下簡稱甲方）與服務執行人 {PROVIDER_NAME}（以下簡稱乙方），茲就廣告投放服務事宜，經雙方同意訂立本合約，條款如下：")
    set_run_font(run_intro, size=11)

    # 定義方案參數
    if payment_opt == "17,000元/月（每月付款）":
        end_dt = start_dt + timedelta(days=30)
        period_text = f"自 {start_dt} 起至 {end_dt} 止，為期 1 個月。本合約屆滿前若雙方未提出終止要求，則自動續約 1 個月，以此類推。"
        fee_text = "新台幣壹萬柒仟元整（NT$17,000）／月。"
        pay_logic = f"應於每月 {pay_day} 日前支付當期費用。"
    else:
        end_dt = start_dt + timedelta(days=90)
        period_text = f"自 {start_dt} 起至 {end_dt} 止，為期 3 個月。續約應於屆滿前 7 日另行協議。"
        fee_text = "新台幣肆萬伍仟元整（NT$45,000）／三個月。"
        pay_logic = f"應於 {pay_dt} 前一次性支付全額費用。"

    # 3. 完整 14 條款
    clauses = [
        ("第一條：合約期間", [period_text]),
        ("第二條：服務內容", [
            "1. 廣告策略規劃與上架執行。",
            "2. 每日監控廣告投放狀況與數據維護。",
            "3. 提供簡易週報（包含數據摘要與下週優化方向）。",
            "4. 視需求提供文案與 Landing Page 優化建議。"
        ]),
        ("第三條：投放平台與費用", [
            "1. 本服務以 Meta（Facebook/Instagram）平台為主。",
            "2. 廣告投放實際消耗之「廣告費」不包含在服務費內，由甲方直接支付予平台方。"
        ]),
        ("第四條：甲方配合義務", [
            "1. 甲方應提供必要之資產存取權限（如粉專管理員、廣告帳號權限）。",
            "2. 因乙方目前帳號限制，甲方同意配合以遠端連線（如 Google Remote Desktop）方式進行必要之後台操作。",
            "3. 甲方應確保廣告素材（圖片、影片）無侵權事宜。"
        ]),
        ("第五條：服務費用與給付方式", [
            f"1. 服務費用：{fee_text}",
            f"2. 支付時間：{pay_logic}",
            f"3. 匯款資訊：{BANK_NAME} ({BANK_CODE}) 帳號：{ACCOUNT_NUMBER}"
        ]),
        ("第六條：稅務說明", [
            "1. 乙方為個人工作室（自然人），本服務費用不開立統一發票。",
            "2. 若甲方需報支費用，請自行依稅法規定開立勞務報酬單，或處理相關代扣繳稅額。"
        ]),
        ("第七條：成果歸屬與智慧財產權", [
            "1. 廣告投放產生之數據與權限歸甲方所有。",
            "2. 乙方所撰寫之文案與投放策略，於合約存續期間授權甲方使用。"
        ]),
        ("第八條：保密義務", [
            "1. 雙方應對合約內容及因履行本合約所獲知之對方商業機密負保密義務。",
            "2. 保密期間自合約簽署日起至終止後兩年止。"
        ]),
        ("第九條：免責聲明與風險承擔", [
            "1. 乙方不保證特定成效指標（如特定 ROAS 或點擊數）。",
            "2. 因平台政策異動、系統故障或不可抗力因素導致廣告中斷，乙方不負賠償責任。"
        ]),
        ("第十條：合約變更", ["本合約之任何修改、變更或補充，均應由雙方以書面（含 LINE/Email）為之。"]),
        ("第十一條：損害賠償", ["如因一方違反本合約導致他方受損，賠償限額以本合約最近一期已支付之服務費用為上限。"]),
        ("第十二條：合約終止", [
            "1. 甲方若欲提前終止，應於 14 日前通知乙方。",
            "2. 已支付之月費或季付優惠費用，於合約啟動後恕不退還。"
        ]),
        ("第十三條：爭議處理與管轄法院", ["如因本合約產生爭議，雙方同意以台灣台北地方法院為第一審管轄法院。"]),
        ("第十四條：其他", ["本合約自簽署日起生效。本合約乙式兩份，由雙方各執一份為憑。"])
    ]

    for title, contents in clauses:
        p_t = doc.add_paragraph()
        r_t = p_t.add_run(title)
        set_run_font(r_t, size=11, bold=True)
        for c in contents:
            p_i = doc.add_paragraph()
            p_i.paragraph_format.left_indent = Cm(0.75)
            r_i = p_i.add_run(c)
            set_run_font(r_i, size=10.5)

    doc.add_paragraph("\n")

    # 4. 簽署欄位
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    
    # 甲方
    c1 = table.cell(0, 0).paragraphs[0]
    run1 = c1.add_run(f"甲方（委託方）：\n{party_a}\n\n簽署人：________________\n\n日期：   年   月   日")
    set_run_font(run1, size=11)
    
    # 乙方
    c2 = table.cell(0, 1).paragraphs[0]
    run2 = c2.add_run(f"乙方（執行方）：\n{PROVIDER_NAME}\n\n簽署人：________________\n\n日期：   年   月   日")
    set_run_font(run2, size=11)

    # 5. 存檔回傳
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# =========================================================
# 4) Sidebar
# =========================================================
if "user" not in st.session_state: st.session_state.user = None 
if "p1_msg" not in st.session_state: st.session_state.p1_msg = None
if "p2_msg" not in st.session_state: st.session_state.p2_msg = None

with st.sidebar:
    st.title("系統入口")
    if st.session_state.user:
        st.success(f"🟢 已登入：{st.session_state.user['name']}")
        with st.expander("🔑 安全與密碼修改"):
            st.warning("⚠️ **安全提示**：本系統採 **SHA-256 雜湊** 加密。請勿使用您的常用密碼（如網銀、Google 密碼）。")
            new_p = st.text_input("設定新密碼", type="password")
            if st.button("確認修改"):
                if len(new_p) < 4: st.error("太短")
                else: update_password(st.session_state.user["row_num"], new_p); st.success("已更新！")
        st.markdown("---")
        if st.button("登出系統"): st.session_state.clear(); st.rerun()
    else:
        mode = st.radio("模式", ["客戶登入", "新客戶建檔"])
        if mode == "新客戶建檔":
            reg_n = st.text_input("客戶名稱")
            reg_e = st.text_input("聯絡信箱 (限 Gmail)")
            if st.button("開始建檔"):
                if not reg_n or not reg_e.endswith("@gmail.com"): st.error("格式錯誤")
                else:
                    row, _ = find_user_row(reg_e)
                    if row: st.error("已註冊。預設密碼 dennis")
                    else: st.session_state.user = {"email": reg_e, "name": reg_n, "role": "new"}; st.rerun()
        else:
            with st.form("login_form"):
                log_e = st.text_input("信箱"); log_p = st.text_input("密碼", type="password")
                if st.form_submit_button("登入"):
                    row, data = find_user_row(log_e)
                    if not row: st.error("找不到此信箱")
                    elif check_password(log_p, str(data.get("password","")).strip() or "dennis"):
                        st.session_state.user = {"email": data["Email"], "name": data["party_a"], "role": "login", "row_num": row, "raw_data": data}
                        st.rerun()
                    else: st.error("密碼錯誤")

# =========================================================
# 5) 渲染
# =========================================================
if not st.session_state.user:
    st.title("📝 廣告投放服務｜合約＋啟動資料收集")
    st.caption("✅ 自動 Word 合約生成 ＋ 雲端資料庫同步")
    st.info("👈 請由左側登入 (預設密碼: dennis) 或建檔。")
    st.stop()

user = st.session_state.user; raw = user.get("raw_data", {})
st.title("📝 廣告投放服務系統")
st.markdown(f"**目前使用者：{user['name']} ({user['email']})**")
st.markdown("---")

nav = st.radio("流程：", ["第一階段｜合約", "第二階段｜啟動前確認"] if user["role"] == "login" else ["第一階段｜合約"], horizontal=True)
if nav == "第一階段｜合約": st.session_state.p2_msg = None
else: st.session_state.p1_msg = None

# 第一階段
if nav == "第一階段｜合約":
    st.header("第一階段｜合約說明與建檔")
    st.info("💡 **操作流程**：詳閱服務內容 -> 選擇方案 -> 生成案件並存檔 -> 下載合約並傳送訊息回傳")
    
    with st.expander("✅ 服務內容與現況提醒 (還原詳細版)", expanded=True):
        st.subheader("固定工作")
        st.markdown("- **廣告上架**\n- **廣告監控 / 維護 / 優化**\n- **簡易週報**（成果摘要、下週優化方向）")
        st.subheader("非固定工作")
        st.markdown("- **廣告文案與素材優化建議**\n- **網頁調整建議**")
        st.info("現況提醒：目前我的 FB 個人帳號被停用，但我仍需要每天監控。我會教你如何每天匯出數據。若需調整後台，我會透過遠端連線操作你的電腦。")
        st.warning("📌 稅務提醒：乙方為自然人，無須開立發票。甲方自行處理勞報或相關稅務。")

    st.subheader("💰 付款方案與日期")
    c1, c2 = st.columns(2)
    # [修改] 移除了 disabled=(user["role"]=="login")，讓登入者可以修改方案與日期
    with c1:
        plan = st.radio("方案選擇：", ["17,000元/月（每月付款）", "45,000元/三個月（一次付款）"], index=0 if raw.get("plan") != "45,000元/三個月（一次付款）" else 1)
        s_date_val = datetime.strptime(raw["start_date"], "%Y-%m-%d").date() if raw.get("start_date") else date.today()+timedelta(days=7)
        s_date = st.date_input("合作啟動日", value=s_date_val)
    with c2:
        p_day = st.slider("每月付款日", 1, 28, int(raw.get("pay_day", 5)) if raw.get("pay_day") else 5) if "每月" in plan else 5
        p_date_val = datetime.strptime(raw["pay_date"], "%Y-%m-%d").date() if raw.get("pay_date") else s_date
        p_date = st.date_input("付款日期", value=p_date_val) if "三個月" in plan else None

    if user["role"] == "new":
        if st.button("🎲 生成案件編號並存檔", type="primary"):
            with st.spinner("建立案件中..."):
                case_id = f"{user['name']}_{datetime.now().strftime('%Y%m%d')}"
                data_dict = {"Email": user["email"], "case_id": case_id, "party_a": user["name"], "plan": plan, "start_date": s_date, "pay_day": p_day, "pay_date": p_date}
                save_phase1_new(data_dict)
                send_email(f"【新案件】{user['name']} 已建檔", f"名稱：{user['name']}\n案件號：{case_id}\n方案：{plan}")
                st.session_state.p1_msg = f"【合約確認】\n案件：{case_id}\n甲方：{user['name']}\n方案：{plan}\n啟動日：{s_date}"
                st.rerun()
    
    if st.session_state.p1_msg:
        st.success("✅ 建檔成功！請複製訊息傳 LINE 給我："); st.code(st.session_state.p1_msg); st.balloons()

    if user["role"] == "login":
        st.info(f"案件編號：{raw.get('case_id')}")

        # [新增] 登入者更新方案按鈕
        if st.button("💾 更新合約方案資料"):
            with st.spinner("更新資料中..."):
                update_phase1(user["row_num"], plan, s_date, p_day, p_date)
                # 更新 session 內的資料，讓介面不需要 F5 就能反映
                st.session_state.user["raw_data"]["plan"] = plan
                st.session_state.user["raw_data"]["start_date"] = str(s_date)
                st.session_state.user["raw_data"]["pay_day"] = p_day
                st.session_state.user["raw_data"]["pay_date"] = str(p_date) if p_date else ""
            st.success("✅ 方案資料已更新！(重新產生合約即可生效)")
            time.sleep(1) # 讓使用者看到成功訊息
            st.rerun()

        if st.button("📝 生成 Word 合約"):
            docx = generate_docx_bytes(user["name"], user["email"], plan, s_date, p_day, p_date, raw.get("case_id"))
            st.download_button("⬇️ 下載 Word 合約 (.docx)", docx, f"合約_{raw.get('case_id')}.docx")

# 第二階段
elif nav == "第二階段｜啟動前確認":
    st.header("第二階段｜啟動資料收集")
    st.info("💡 **操作流程**：確認資產現況 -> 填寫行銷情報 -> 更新並傳送訊息回傳")
    
    if BM_TUTORIAL_URL:
        with st.expander("📺 [教學影片] 如何設定企業管理平台 (BM)？", expanded=True):
            st.video(BM_TUTORIAL_URL)

    def b(k): return str(raw.get(k, "FALSE")).upper() == "TRUE"
    def s(k): return raw.get(k, "")

    st.subheader("✅ 確認事項 (照實勾選)")
    col1, col2 = st.columns(2)
    with col1:
        ad = st.checkbox("廣告帳號已開啟", value=b("chk_ad_account"))
        px = st.checkbox("像素事件已埋放", value=b("chk_pixel"))
    with col2:
        fp = st.checkbox("粉專已建立", value=b("chk_fanpage"))
        bm = st.checkbox("企業管理平台已建立", value=b("chk_bm"))
    
    rem = st.checkbox("已完成 Google 遠端桌面設定 (提醒)", value=b("chk_remote"))
    st.caption(f"[🔗 遠端教學連結]({REMOTE_SUPPORT_URL})")
    cre = st.checkbox("已前往素材系統上傳素材", value=b("chk_creatives"))
    st.caption(f"[🔗 素材系統連結]({CREATIVES_UPLOAD_URL})")

    st.markdown("---")
    st.subheader("🧾 須提供事項與行銷情報")
    fp_u = st.text_input("粉專網址", value=s("fanpage_url"))
    ld_u = st.text_input("廣告導向頁 (Landing Page)", value=s("landing_url"))
    
    st.markdown("**競爭對手粉專**")
    cp1 = st.text_input("競品 1", value=s("comp1"))
    cp2 = st.text_input("競品 2", value=s("comp2"))
    cp3 = st.text_input("競品 3", value=s("comp3"))
    
    st.markdown("**定位與痛點**")
    who = st.text_area("解決誰的問題？", value=s("who_problem"))
    what = st.text_area("要解決什麼問題？", value=s("what_problem"))
    how = st.text_area("如何解決？", value=s("how_solve"))
    bud = st.text_input("第一個月預算", value=s("budget"))

    if st.button("💾 更新資料並通知", type="primary"):
        with st.spinner("同步雲端資料中..."):
            p2_data = {
                "chk_ad_account": ad, "chk_pixel": px, "chk_fanpage": fp, "chk_bm": bm, "chk_remote": rem, "chk_creatives": cre,
                "fanpage_url": fp_u, "landing_url": ld_u, "comp1": cp1, "comp2": cp2, "comp3": cp3,
                "who_problem": who, "what_problem": what, "how_solve": how, "budget": bud
            }
            update_phase2(user["row_num"], p2_data)
            send_email(f"【更新】{user['name']} 啟動資料", f"客戶 {user['name']} 已更新啟動資料。")
            st.session_state.p2_msg = f"""【資料更新】
案件編號：{raw.get('case_id')}
遠端桌面：{'OK' if rem else '未完成'}
素材上傳：{'OK' if cre else '未完成'}
預算：{bud}
競品：{cp1}, {cp2}, {cp3}"""
            st.rerun()

    if st.session_state.p2_msg:
        st.success("✅ 更新成功！請複製以下訊息傳 LINE 給我："); st.code(st.session_state.p2_msg); st.balloons()
